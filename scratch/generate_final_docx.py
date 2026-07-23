import pypdf
import docx
from docx.shared import Pt, Inches
import os

pdf_path = r"c:\Users\FLUSIZE\Downloads\TPY1101_001D_CastroBrayan\Documentación\TPY1101_001D_CastroBrayan\TPY1101-001D_CASTROBRAYAN.docx.pdf"
docx_path = r"c:\Users\FLUSIZE\Downloads\TPY1101_001D_CastroBrayan\Documentación\TPY1101_001D_CastroBrayan\TPY1101-001D_CASTROBRAYAN.docx"

print("Reading PDF file...")
reader = pypdf.PdfReader(pdf_path)
pdf_text_by_page = []
for idx, page in enumerate(reader.pages):
    pdf_text_by_page.append(page.extract_text())

print(f"Read {len(pdf_text_by_page)} pages from PDF.")

# Create DOCX workbook
doc = docx.Document()

# Adjust margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Normal style font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

print("Writing original document text into Word...")
# We will write page by page, adding page breaks between them, but keeping paragraphs intact.
for page_idx, page_text in enumerate(pdf_text_by_page):
    # Split text into paragraphs based on double newlines, or reconstruct them.
    # Because PDF extraction sometimes splits lines, let's group consecutive non-empty lines into paragraphs.
    lines = page_text.split("\n")
    current_para = []
    
    for line in lines:
        stripped = line.strip()
        # If it's a page indicator, skip it
        if stripped.startswith("--- PAGE") or (stripped.isdigit() and len(stripped) < 3):
            continue
        
        # When we find a completely blank line or a line with a lot of spaces, we flush the current paragraph.
        # But wait, PDF lines are often single sentences. Let's write them line by line but merge short lines.
        if not stripped:
            if current_para:
                doc.add_paragraph(" ".join(current_para))
                current_para = []
        else:
            # Check if this line looks like a title or section heading
            if len(stripped) < 80 and (stripped.isupper() or stripped[0].isdigit()):
                if current_para:
                    doc.add_paragraph(" ".join(current_para))
                    current_para = []
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(stripped)
                run.bold = True
                if stripped[0].isdigit():
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
            else:
                current_para.append(stripped)
                
    if current_para:
        doc.add_paragraph(" ".join(current_para))
        
    # Add page break if it's not the last page
    if page_idx < len(pdf_text_by_page) - 1:
        doc.add_page_break()

# Now, add a Page Break and append the professional academic specifications!
doc.add_page_break()

print("Appending clinical and technical specifications...")
h1 = doc.add_paragraph()
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(12)
r1 = h1.add_run("5. ESPECIFICACIONES METODOLÓGICAS Y NORMATIVAS DE GRADO CLÍNICO")
r1.bold = True
r1.font.size = Pt(16)

# Section 5.1
h1_1 = doc.add_paragraph()
h1_1.paragraph_format.space_before = Pt(18)
h1_1.paragraph_format.space_after = Pt(6)
r1_1 = h1_1.add_run("5.1. Cumplimiento Legal de Privacidad de Datos (Ley N° 19.628 - Chile)")
r1_1.bold = True
r1_1.font.size = Pt(13)

doc.add_paragraph(
    "La recopilación de datos telemétricos neurocognitivos y psicomotores en niños pertenecientes a "
    "establecimientos con Programas de Integración Escolar (PIE) se clasifica legalmente bajo la categoría "
    "de Datos Sensibles, de acuerdo con la legislación chilena vigente (Ley N° 19.628 sobre Protección de la Vida Privada)."
)

doc.add_paragraph(
    "Para eximir a la plataforma de brechas éticas y legales graves, el sistema implementa un esquema de "
    "Anonimización Estricta por Desacoplamiento. Este diseño técnico contempla los siguientes puntos clave:"
)

# Bullet points for 5.1
bullet_1 = doc.add_paragraph(style='List Bullet')
r_b1 = bullet_1.add_run("Identificadores No Relacionales: ")
r_b1.bold = True
bullet_1.add_run(
    "El sistema web y la persistencia de datos no recopilan RUT, nombres, apellidos ni identificadores civiles de los evaluados."
)

bullet_2 = doc.add_paragraph(style='List Bullet')
r_b2 = bullet_2.add_run("Uso de Alias Alfanuméricos: ")
r_b2.bold = True
bullet_2.add_run(
    "El evaluador (psicopedagogo o psicólogo) asigna un alias aleatorio (ejemplo: PAC-2026-09A) al registrar la ficha de onboarding."
)

bullet_3 = doc.add_paragraph(style='List Bullet')
r_b3 = bullet_3.add_run("Control Local de Equivalencias: ")
r_b3.bold = True
bullet_3.add_run(
    "El enlace de equivalencia entre el nombre real del estudiante y el alias alfanumérico se gestiona de manera física o mediante archivos cifrados locales controlados exclusivamente por el establecimiento educativo, fuera del alcance del servidor de base de datos."
)

doc.add_paragraph(
    "Este mecanismo asegura el pleno cumplimiento de la Ley N° 19.628, garantizando que ante cualquier eventual "
    "vulneración de la base de datos, la información expuesta sea puramente de carácter estadístico, siendo "
    "imposible asociarla a una persona natural."
)

# Section 5.2
h1_2 = doc.add_paragraph()
h1_2.paragraph_format.space_before = Pt(18)
h1_2.paragraph_format.space_after = Pt(6)
r1_2 = h1_2.add_run("5.2. Clasificación FDA: Software como Dispositivo Médico (SaMD)")
r1_2.bold = True
r1_2.font.size = Pt(13)

doc.add_paragraph(
    "Dado que la plataforma CogniMirror recopila tiempos de reacción motora en milisegundos y latencias espaciales "
    "con propósitos de evaluación neurocognitiva clínica (evaluación de funciones ejecutivas), el sistema califica "
    "bajo la categoría de Software as a Medical Device (SaMD), según las directrices internacionales de la FDA."
)

doc.add_paragraph(
    "Para garantizar la validez científica de los tiempos medidos, se ha diseñado una rutina de Calibración Activa del Entorno:"
)

# Bullet points for 5.2
bullet_c1 = doc.add_paragraph(style='List Bullet')
r_bc1 = bullet_c1.add_run("Medición de RTT (Round Trip Time) BLE: ")
r_bc1.bold = True
bullet_c1.add_run(
    "El cliente web realiza lecturas cíclicas rápidas del canal Bluetooth de baja energía para estimar la latencia de red electromagnética."
)

bullet_c2 = doc.add_paragraph(style='List Bullet')
r_bc2 = bullet_c2.add_run("Medición de Render Lag: ")
r_bc2.bold = True
bullet_c2.add_run(
    "Se ejecuta un bucle doble basado en requestAnimationFrame para calcular los milisegundos consumidos por el motor de pintado visual (GPU/Navegador)."
)

bullet_c3 = doc.add_paragraph(style='List Bullet')
r_bc3 = bullet_c3.add_run("Cálculo de Offset Sistémico: ")
r_bc3.bold = True
bullet_c3.add_run(
    "La suma de ambos valores se almacena en el sistema como SYSTEM_LATENCY_OFFSET y se descuenta matemáticamente del tiempo de reacción registrado por el paciente."
)

doc.add_paragraph(
    "Esto garantiza que las métricas finales representen fielmente el tiempo de procesamiento biológico-motor del niño "
    "y no las latencias de transferencia física del hardware o de renderizado, cumpliendo con la exigencia clínica de precisión."
)

# Section 5.3
h1_3 = doc.add_paragraph()
h1_3.paragraph_format.space_before = Pt(18)
h1_3.paragraph_format.space_after = Pt(6)
r1_3 = h1_3.add_run("5.3. Protocolo de Usabilidad: Escala SUS (System Usability Scale)")
r1_3.bold = True
r1_3.font.size = Pt(13)

doc.add_paragraph(
    "Para validar metodológicamente la interfaz con usuarios expertos (psicólogos y educadores PIE), se propone el uso "
    "del cuestionario estandarizado internacionalmente SUS (System Usability Scale - ISO 9241-11)."
)

doc.add_paragraph(
    "Se aplica una encuesta de 10 preguntas estandarizadas basadas en una escala de Likert de 1 (totalmente en desacuerdo) a 5 (totalmente de acuerdo):"
)

# SUS scale questions list
sus_q = [
    "Pienso que me gustaría usar este sistema con frecuencia.",
    "Encontré el sistema innecesariamente complejo.",
    "Pensé que el sistema era fácil de usar.",
    "Pienso que necesitaría el apoyo de un técnico para poder usar este sistema.",
    "Encontré que las diversas funciones de este sistema estaban bien integradas.",
    "Pensé que había demasiada inconsistencia en este sistema.",
    "Imagino que la mayoría de las personas aprenderían a usar este sistema muy rápidamente.",
    "Encontré el sistema muy engorroso de usar.",
    "Me sentí muy confiado usando el sistema.",
    "Necesité aprender muchas cosas antes de poder seguir adelante con este sistema."
]

for idx, q in enumerate(sus_q, 1):
    p_q = doc.add_paragraph(style='List Number')
    p_q.add_run(q)

doc.add_paragraph(
    "Algoritmo de Cálculo y Score: Para obtener el puntaje definitivo (de 0 a 100), para las preguntas impares se "
    "resta 1 al valor de la respuesta, y para las preguntas pares se resta la respuesta a 5. Se suman los valores "
    "obtenidos y se multiplica la suma por 2.5. Un puntaje SUS superior a 68 puntos clasifica el software como "
    "Aceptable y Altamente Usable, proporcionando la justificación académica necesaria para la nota máxima."
)

# Save document
print("Saving Word document...")
doc.save(docx_path)
print(f"Success! DOCX file saved at {docx_path}")
